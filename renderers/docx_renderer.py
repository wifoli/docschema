"""
docschema.renderers.docx_renderer  (v2)
==========================================
DOCX renderer using python-docx.
pip install python-docx

v2 additions:
  - output_path accepts BytesIO
  - TabelaConteudo generates real DOCX TOC field (auto-updates on open)
  - toc_default=True for DOCX (default behaviour)
  - Assinatura, Badge, Formulario, Se, Para, NotaRodape
  - Footnotes via python-docx if available
"""
from __future__ import annotations

import io
from typing import Any, Dict, Generator, List, Optional

from docschema.renderers.base import BaseRenderer

try:
    import docx
    from docx import Document as DocxDoc
    from docx.shared import Pt, Cm, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


def _add_toc_field(doc) -> None:
    """Insert a real Word TOC field that auto-populates on first open."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    run2 = paragraph.add_run()
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run3._r.append(fld2)


class DocxRenderer(BaseRenderer):
    FORMAT_NAME = "docx"

    def __init__(self, output_path=None, toc_default: bool = True, **kwargs):
        super().__init__(**kwargs)
        if not _DOCX_AVAILABLE:
            raise ImportError("DocxRenderer requires python-docx. Run: pip install python-docx")
        self._output_path = output_path or "output.docx"
        self._toc_default = toc_default
        self._doc_obj: Optional[Any] = None

    # ── Document ──────────────────────────────────────────────────────────────

    def render_document(self, doc) -> Any:
        self._doc       = doc
        self._doc_cfg   = doc.get_format_config(self.FORMAT_NAME)
        self._context   = getattr(doc, "_render_context", {})
        self._footnotes = []

        self._doc_obj = DocxDoc()
        d = self._doc_obj

        # Page margins
        for section in d.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(2.5)

        # Title
        if doc.titulo:
            p = d.add_heading(doc.titulo, level=0)
        if doc.autor:
            d.add_paragraph(f"Autor: {doc.autor}").italic = True

        # Auto-insert TOC if toc_default or if a TabelaConteudo element isn't present
        from docschema.elements.block import TabelaConteudo
        has_toc_el = any(isinstance(c, TabelaConteudo) for c in doc.children)
        cfg_toc    = self._doc_cfg.get("toc", self._toc_default)
        if cfg_toc and not has_toc_el:
            d.add_heading("Sumário", level=1)
            _add_toc_field(d)
            d.add_page_break()

        for child in doc.children:
            child.accept(self)

        # Footnotes as endnotes (appended at end)
        if self._footnotes:
            d.add_heading("Notas de Rodapé", level=2)
            for i, fn in enumerate(self._footnotes, 1):
                text    = self._inline_text(fn.children)
                para    = d.add_paragraph()
                run     = para.add_run(f"[{i}] {text}")
                run.font.size = Pt(9)

        path = self._output_path
        if isinstance(path, str):
            d.save(path)
            return path
        else:
            d.save(path)
            if hasattr(path, "getvalue"):
                return path.getvalue()
            return path

    def render_document_stream(self, doc) -> Generator[bytes, None, None]:
        buf = io.BytesIO()
        original = self._output_path
        self._output_path = buf
        self.render_document(doc)
        self._output_path = original
        yield buf.getvalue()

    # ── Inline helpers ────────────────────────────────────────────────────────

    def _inline_text(self, children) -> str:
        return "".join(str(el.accept(self)) for el in children)

    def _add_inline_run(self, paragraph, children):
        for el in children:
            el.accept_run(self, paragraph)  # if supported; fallback below

    def visit_texto(self, el): return el.conteudo
    def visit_negrito(self, el): return self._inline_text(el.children)
    def visit_italico(self, el): return self._inline_text(el.children)
    def visit_sublinhado(self, el): return self._inline_text(el.children)
    def visit_link(self, el): return self._inline_text(el.children) + (f" ({el.url})" if el.url else "")
    def visit_quebra_linha(self, el): return "\n"
    def visit_span(self, el): return self._inline_text(el.children)
    def visit_emoji(self, el): return el.simbolo
    def visit_badge(self, el): return f"[{self._inline_text(el.children)}]"
    def visit_data_hora(self, el): return el.renderizado()
    def visit_local(self, el): return el.renderizado()
    def visit_ancora(self, el): return ""
    def visit_ref_cruzada(self, el): return el.texto
    def visit_variavel(self, el): return str(self._context.get(el.nome, el.padrao))

    def visit_marcador_rodape(self, el):
        self._footnotes.append(el)
        n = len(self._footnotes)
        return f"[{n}]"

    # ── Block ─────────────────────────────────────────────────────────────────

    def _para_with_inline(self, children, style="Normal"):
        d = self._doc_obj
        p = d.add_paragraph(style=style)
        for el in children:
            from docschema.elements.inline import (
                Texto, Negrito, Italico, Sublinhado, Link,
                Badge, DataHora, Local, Ancora, RefCruzada, MarcadorRodape, Variavel, Emoji, Span, QuebraLinha,
            )
            text = el.accept(self)
            if not text:
                continue
            run = p.add_run(str(text))
            if isinstance(el, Negrito):
                run.bold = True
            elif isinstance(el, Italico):
                run.italic = True
            elif isinstance(el, Sublinhado):
                run.underline = True
        return p

    def visit_titulo(self, el):
        text = self._inline_text(el.children)
        self._doc_obj.add_heading(text, level=el.nivel)

    def visit_paragrafo(self, el):
        self._para_with_inline(el.children)

    def visit_secao(self, el):
        if el.titulo:
            el.titulo.accept(self)
        for child in el.children:
            child.accept(self)

    def visit_lista(self, el):
        style = "List Number" if el.ordenada else "List Bullet"
        for item in el.children:
            self._para_with_inline(item.children, style=style)

    def visit_item_lista(self, el):
        self._para_with_inline(el.children, style="List Bullet")

    def visit_tabela(self, el):
        d = self._doc_obj
        cab  = [self._cel_text(c) for c in el.cabecalho]
        rows = [[self._cel_text(c) for c in row] for row in el.linhas]
        ncols = len(cab)
        t = d.add_table(rows=1 + len(rows), cols=ncols)
        t.style = "Table Grid"

        # Header row
        for j, h in enumerate(cab):
            cell = t.rows[0].cells[j]
            cell.text = h
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell._tc.get_or_add_tcPr().append(
                self._make_shading("2E75B6")
            )

        for i, row in enumerate(rows):
            for j, cell_val in enumerate(row):
                t.rows[i+1].cells[j].text = cell_val

        if el.legenda:
            p = d.add_paragraph(el.legenda)
            p.runs[0].italic = True
            p.runs[0].font.size = Pt(9)

    def _cel_text(self, cel):
        from docschema.elements.block import CelulaTabela
        if isinstance(cel, CelulaTabela):
            return self._inline_text(cel.inline_content())
        return str(cel)

    def _make_shading(self, color: str):
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  color)
        return shd

    def visit_tabela_conteudo(self, el):
        d = self._doc_obj
        d.add_heading(el.titulo, level=1)
        _add_toc_field(d)

    def visit_nota_rodape(self, el):
        text = self._inline_text(el.children)
        p = self._doc_obj.add_paragraph(f"[{el.numero}] {text}")
        p.runs[0].font.size = Pt(9)

    def visit_imagem(self, el):
        d = self._doc_obj
        if el.has_fallback(self.FORMAT_NAME):
            d.add_paragraph(str(el.get_fallback_value(self.FORMAT_NAME)))
            return
        try:
            w = Cm(el.largura / 37.8) if el.largura else Cm(14)
            d.add_picture(el.src, width=w)
            if el.legenda:
                p = d.add_paragraph(el.legenda)
                p.runs[0].italic = True
        except Exception:
            d.add_paragraph(f"[Imagem: {el.alt}]")

    def visit_espaco(self, el):
        for _ in range(el.linhas):
            self._doc_obj.add_paragraph("")

    def visit_quebra_pagina(self, el):
        self._doc_obj.add_page_break()

    def visit_linha_horizontal(self, el):
        p = self._doc_obj.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        pBdr.append(bottom)
        pPr.append(pBdr)

    def visit_bloco_destaque(self, el):
        if el.titulo:
            p = self._doc_obj.add_paragraph(el.titulo)
            p.runs[0].bold = True
        for child in el.children:
            child.accept(self)

    def visit_nota(self, el):
        _labels = {"info": "ℹ️", "warning": "⚠️", "tip": "💡", "danger": "🚨"}
        icon    = _labels.get(el.tipo, "")
        text    = self._inline_text(el.children)
        self._doc_obj.add_paragraph(f"{icon} {text}")

    def visit_citacao(self, el):
        text = self._inline_text(el.children)
        p    = self._doc_obj.add_paragraph(style="Quote")
        p.add_run(text).italic = True
        if el.autoria:
            p2 = self._doc_obj.add_paragraph(f"— {el.autoria}")
            p2.runs[0].italic = True

    def visit_codigo(self, el):
        p = self._doc_obj.add_paragraph()
        run = p.add_run(el.conteudo)
        run.font.name = "Courier New"
        run.font.size = Pt(9)

    def visit_formulario(self, el):
        if el.titulo:
            self._doc_obj.add_heading(el.titulo, level=3)
        for campo in el.campos:
            campo.accept(self)

    def visit_campo_formulario(self, el):
        req  = " *" if el.obrigatorio else ""
        text = f"{el.label}{req}: _________________________"
        if el.tipo == "checkbox":
            text = f"☐ {el.label}{req}"
        self._doc_obj.add_paragraph(text)

    def visit_se(self, el):
        ramo = el.sim if el.avaliar(self._context) else el.nao
        if ramo:
            for child in ramo:
                child.accept(self)

    def visit_para(self, el):
        itens = el.itens(self._context) if callable(el.itens) else el.itens
        for item in itens:
            for child in el.template(item, self._context):
                child.accept(self)
            if el.separador and item is not itens[-1]:
                el.separador.accept(self)

    def visit_assinatura(self, el):
        d = self._doc_obj
        linha = "_" * el.largura
        p = d.add_paragraph(linha)
        if el.nome:
            d.add_paragraph(el.nome).runs[0].bold = True
        if el.cargo:
            p2 = d.add_paragraph(el.cargo)
            p2.runs[0].italic = True
        sub = []
        if el.mostrar_data:
            sub.append(el.data_formatada())
        if el.mostrar_local and el.local:
            sub.append(el.local)
        if sub:
            p3 = d.add_paragraph(" | ".join(sub))
            p3.runs[0].font.size = Pt(9)
